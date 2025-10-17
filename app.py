import streamlit as st
import pandas as pd
from docx import Document
from datetime import datetime, date, timedelta
import os
import re  # Para extração do valor numérico da idade
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from io import BytesIO

# Caminho para a imagem do emblema (o logo do INS foi removido)
EMBLEM_PATH = "Emblem_of_Mozambique.svg.png"

# Configuração inicial do Streamlit (antes de qualquer saída)
st.set_page_config(
    page_title="Gerador de Relatórios - Vigilância das Infecções Respiratórias",
    page_icon="🦠",
    layout="wide"
)

# CONSTANTE: data/hora do sistema para rodapé
CURRENT_DATE = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def extrair_valor_idade(idade_str):
    """
    Extrai o valor numérico da idade considerando:
      'a' para anos, 'm' para meses (dividindo por 12) e 'd' para dias (dividindo por 365).
    Retorna None se o padrão não for reconhecido.
    """
    if isinstance(idade_str, str):
        match = re.match(r'(\d+)([amd])', idade_str.lower())
        if match:
            valor = int(match.group(1))
            unidade = match.group(2)
            if unidade == 'a':
                return valor
            elif unidade == 'm':
                return valor / 12
            elif unidade == 'd':
                return valor / 365
    return None

def classificar_influenza_subtipos(row):
    """
    Verifica os subtipos de Influenza nas colunas:
    InfA, Apdm, H1pdm, H3, H5, H5a, H5b, H7, InfB, Vic, Yam.
    Para cada coluna, tenta converter o valor para float; se o Ct for menor que 40,
    considera positivo para aquele subtipo.
    Se encontrar algum, retorna "POSITIVO: [subtipo(s)]", senão retorna "NEGATIVO".
    """
    columns_subtipos = {
        "InfA":  "A",
        "Apdm":  "A(H1pdm)",
        "H1pdm": "A(H1pdm)",
        "H3":    "A(H3N2)",
        "H5":    "A(H5)",
        "H5a":   "A(H5a)",
        "H5b":   "A(H5b)",
        "H7":    "A(H7)",
        "InfB":  "B",
        "Vic":   "B(Victoria)",
        "Yam":   "B(Yamagata)"
    }
    THRESHOLD_CT = 40.0
    found_subtypes = []
    for col, label in columns_subtipos.items():
        val = row.get(col, None)
        try:
            ct_value = float(val)
            if ct_value < THRESHOLD_CT:
                found_subtypes.append(label)
        except (ValueError, TypeError):
            pass
    if found_subtypes:
        return "POSITIVO: " + ", ".join(found_subtypes)
    else:
        return "NEGATIVO"

def carregar_dados(uploaded_file):
    """
    Lê e processa os dados do arquivo, utilizando a coluna "Código do Site"
    e garantindo a normalização das colunas obrigatórias.
    """
    try:
        if uploaded_file.name.endswith('.csv'):
            df = pd.read_csv(uploaded_file)
        elif uploaded_file.name.endswith(('.xls', '.xlsx')):
            df = pd.read_excel(uploaded_file)
        else:
            raise ValueError("Formato de ficheiro não suportado. Use .csv ou .xlsx")
        
        df.columns = df.columns.str.strip().str.replace("  ", " ", regex=False)
        
        colunas_obrigatorias = [
            "Código do Site", "Sexo", "Idade", "Residência/Bairro",
            "Data da Colheita", "Data de entrada", "Resultado RSV"
        ]
        colunas_faltantes = [col for col in colunas_obrigatorias if col not in df.columns]
        if colunas_faltantes:
            raise ValueError(f"Colunas obrigatórias faltando: {', '.join(colunas_faltantes)}")
        
        # Conversão de colunas principais em datetime
        df['Data da Colheita'] = pd.to_datetime(df['Data da Colheita'], errors='coerce')
        df['Data de entrada'] = pd.to_datetime(df['Data de entrada'], errors='coerce')

        # ✅ NOVO: converte colunas opcionais de testagem, caso existam
        for col in ["Data de Testagem SARS", "Data da Testagem FLU", "Data da Testagem RSV", "Data da Testagem"]:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], errors='coerce')

        df['Idade_Num'] = df['Idade'].apply(extrair_valor_idade)
        
        # Cria a coluna Influenza utilizando os subtipos (detalhamento somente para resumo)
        df['Influenza'] = df.apply(classificar_influenza_subtipos, axis=1)
        
        # --- NOVO: flags de Influenza A e B por CT < 40 ---
        def ct_pos(v):
            try:
                return float(v) < 40
            except:
                return False

        fluA_cols = ["InfA", "Apdm", "H1pdm", "H3", "H5", "H5a", "H5b", "H7"]
        fluB_cols = ["InfB", "Vic", "Yam"]
        for c in fluA_cols + fluB_cols:
            if c not in df.columns:
                df[c] = None

        df["FluA_Pos"] = df[fluA_cols].apply(lambda r: any(ct_pos(x) for x in r), axis=1)
        df["FluB_Pos"] = df[fluB_cols].apply(lambda r: any(ct_pos(x) for x in r), axis=1)

        # Trata RSV
        df['Resultado RSV'] = df['Resultado RSV'].fillna("-").astype(str).str.upper()
        
        # Para SARS-CoV-2, utiliza a primeira coluna encontrada dentre as possíveis
        colunas_sars = ["Resultado SARS", "Resultado Sars-Cov-2", "Resultado  SARS"]
        resultado_sars_col = next((col for col in colunas_sars if col in df.columns), None)
        if not resultado_sars_col:
            raise ValueError("Coluna de resultado SARS-CoV-2 não encontrada.")
        df[resultado_sars_col] = df[resultado_sars_col].fillna("-").astype(str).str.upper()
        
        df_limpo = pd.DataFrame({
            "Código": df["Código do Site"].astype(str).str.strip(),
            "Sexo": df["Sexo"].astype(str).str.upper(),
            "Idade": df["Idade"].astype(str),
            "Residência/Bairro": df["Residência/Bairro"].astype(str).fillna("Não especificado"),
            "Data da Colheita": df["Data da Colheita"],
            "Data de entrada": df["Data de entrada"],
            "Tipo de Amostra": "Nasofaríngeo",
            # Na tabela, para Influenza, usaremos apenas "POSITIVO"/"NEGATIVO"
            "Influenza": df["Influenza"],
            "RSV": df["Resultado RSV"],
            "SARS-CoV-2": df[resultado_sars_col]
        })

        # ✅ Preservar, se existirem
        for col in ["Data de Testagem SARS", "Data da Testagem FLU", "Data da Testagem RSV", "Data da Testagem"]:
            if col in df.columns:
                df_limpo[col] = df[col]
        df_limpo["FluA_Pos"] = df["FluA_Pos"]
        df_limpo["FluB_Pos"] = df["FluB_Pos"]

        if df_limpo.empty:
            raise ValueError("Nenhum dado válido encontrado após processamento.")
        return df_limpo

    except Exception as e:
        st.error(f"Erro ao processar arquivo: {str(e)}")
        return None

def calcular_resumo(df):
    """Calcula estatísticas globais para os patógenos."""
    total = len(df)
    if total == 0:
        return {
            "total": 0,
            "pos_influenza": 0, "pos_influenza_perc": 0,
            "pos_sars": 0, "pos_sars_perc": 0,
            "pos_rsv": 0, "pos_rsv_perc": 0
        }
    pos_influenza = df["Influenza"].str.upper().str.contains("POSITIVO").sum()
    pos_sars = (df["SARS-CoV-2"].astype(str).str.strip().str.upper() == "POSITIVO").sum()
    pos_rsv = (df["RSV"].astype(str).str.strip().str.upper() == "POSITIVO").sum()
    return {
        "total": total,
        "pos_influenza": pos_influenza,
        "pos_influenza_perc": round(pos_influenza / total * 100, 2),
        "pos_sars": pos_sars,
        "pos_sars_perc": round(pos_sars / total * 100, 2),
        "pos_rsv": pos_rsv,
        "pos_rsv_perc": round(pos_rsv / total * 100, 2)
    }

def gerar_resumo_dinamico(df_atual, df_anterior, periodo_atual_str, periodo_anterior_str):
    """
    Gera um resumo dinâmico:
      - Para cada unidade sanitária do período atual, informa o número total de amostras e, para Influenza,
        exibe o número de positivos por cada subtipo (usando os detalhes); para SARS-CoV-2 e RSV, exibe o número de positivos.
        Utiliza termos singulares quando apropriado.
      - Em seguida, compara as taxas globais com o período anterior.
    """
    total_current = len(df_atual)
    resumo = f"No período entre {periodo_atual_str}, foram testadas {total_current} " \
             f"{'amostra' if total_current == 1 else 'amostras'}.\n"
    resumo += "Resumo por unidade sanitária:\n"
    
    unidades = {
       "IRAS1": "HCM pediatria",
       "IRAS2": "HGM pediatria",
       "IRAS3": "Centro de saude de Mavalane",
       "IRAS4": "Centro de saude de Marracuene",
       "IRAS5": "Hospital Centra de Maputo Adultos",
       "IRAS6": "Hospital Geral de Mavalane Adulto",
       "IDS": "CSZ"
    }
    lines = []
    for code, nome in unidades.items():
        df_unit = df_atual[df_atual["Código"].str.startswith(code, na=False)]
        if not df_unit.empty:
            n_total = len(df_unit)
            # Para Influenza, conta positivos por subtipo usando os detalhes (que estão no formato "POSITIVO: A(H3N2), B(Victoria)" )
            influ_counts = {}
            for val in df_unit["Influenza"]:
                if "POSITIVO:" in str(val).upper():
                    subtypes_str = str(val).split(":", 1)[1].strip()
                    if subtypes_str:
                        subs = [s.strip() for s in subtypes_str.split(",")]
                        for s in subs:
                            influ_counts[s] = influ_counts.get(s, 0) + 1
            if influ_counts:
                influ_detail = ", ".join([f"{subtype}: {count} " + ("positiva" if count == 1 else "positivas") 
                                           for subtype, count in influ_counts.items()])
            else:
                influ_detail = "0"
            n_sars = (df_unit["SARS-CoV-2"].astype(str).str.strip().str.upper() == "POSITIVO").sum()
            n_rsv = (df_unit["RSV"].astype(str).str.strip().str.upper() == "POSITIVO").sum()
            sars_text = f"{n_sars} " + ("positiva" if n_sars == 1 else "positivas")
            rsv_text = f"{n_rsv} " + ("positiva" if n_rsv == 1 else "positivas")
            line = (f"{n_total} " + ("amostra" if n_total == 1 else "amostras") +
                    f" provenientes de {nome}: Influenza: {influ_detail}; "
                    f"SARS-CoV-2: {sars_text}; RSV: {rsv_text}.")
            lines.append(line)
    if lines:
        resumo += "\n".join(lines) + "\n"
    else:
        resumo += "Nenhuma unidade sanitária possui dados específicos para o período.\n"
    
    def calc_taxas(df):
        total = len(df)
        if total == 0:
            return {"Influenza": 0, "SARS-CoV-2": 0, "RSV": 0}
        return {
            "Influenza": round(100 * (df["Influenza"].astype(str).str.strip().str.upper().str.contains("POSITIVO")).sum() / total, 2),
            "SARS-CoV-2": round(100 * (df["SARS-CoV-2"].astype(str).str.strip().str.upper() == "POSITIVO").sum() / total, 2),
            "RSV": round(100 * (df["RSV"].astype(str).str.strip().str.upper() == "POSITIVO").sum() / total, 2)
        }
    rates_curr = calc_taxas(df_atual)
    rates_prev = calc_taxas(df_anterior)
    
    resumo += (f"\nComparando com a semana anterior ({periodo_anterior_str}):\n"
               f"Influenza: {rates_prev['Influenza']}% na semana anterior, {rates_curr['Influenza']}% na presente semana.\n"
               f"SARS-CoV-2: {rates_prev['SARS-CoV-2']}% na semana anterior, {rates_curr['SARS-CoV-2']}% na presente semana.\n"
               f"RSV: {rates_prev['RSV']}% na semana anterior, {rates_curr['RSV']}% na presente semana.")
    
    return resumo

def criar_tabelas_unidades_sanitarias(doc, df):
    """
    Para cada unidade sanitária com registros (usando str.startswith na coluna "Código"),
    gera uma tabela (centralizada com textos centralizados) contendo:
      Ordem, Código, Sexo, Idade, Residência/Bairro, Data de colheita,
      Tipo de Amostra, Influenza, RSV e SARS-CoV-2.
    Na tabela, a coluna de Influenza exibe apenas "POSITIVO" ou "NEGATIVO".
    Aplica formatação condicional (negrito e vermelho) aos resultados positivos nas colunas de resultado.
    Se nenhuma unidade tiver dados, exibe uma mensagem.
    """
    unidades = {
       "IRAS1": "HCM pediatria",
       "IRAS2": "HGM pediatria",
       "IRAS3": "Centro de saude de Mavalane",
       "IRAS4": "Centro de saude de Marracuene",
       "IRAS5": "Hospital Centra de Maputo Adultos",
       "IRAS6": "Hospital Geral de Mavalane Adulto",
       "IDS": "CSZ"
    }
    ordem_unidades = list(unidades.keys())
    tabela_numero = 1
    count_tables = 0

    for cod_site in ordem_unidades:
        df_site = df[df["Código"].str.startswith(cod_site, na=False)]
        if not df_site.empty:
            count_tables += 1
            nome_unidade = unidades.get(cod_site, cod_site)
            p = doc.add_paragraph(f"Tabela {tabela_numero}. Resultado de testagem das amostras provenientes de {nome_unidade}",
                                  style='Heading 2')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            table = doc.add_table(rows=1, cols=10)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Ordem"
            hdr_cells[1].text = "Código"
            hdr_cells[2].text = "Sexo"
            hdr_cells[3].text = "Idade"
            hdr_cells[4].text = "Residência/Bairro"
            hdr_cells[5].text = "Data de colheita"
            hdr_cells[6].text = "Tipo de Amostra"
            hdr_cells[7].text = "Influenza"
            hdr_cells[8].text = "RSV"
            hdr_cells[9].text = "SARS-CoV-2"

            for idx, row in enumerate(df_site.itertuples(index=False), start=1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = str(row[0])
                row_cells[2].text = str(row[1])
                row_cells[3].text = str(row[2])
                row_cells[4].text = str(row[3])
                row_cells[5].text = row[4].strftime('%d/%m/%Y') if pd.notnull(row[4]) else ""
                row_cells[6].text = str(row[5])
                # Exibe apenas "POSITIVO" ou "NEGATIVO" para Influenza
                influenza_val = str(row[6])
                if "POSITIVO" in influenza_val.upper():
                    row_cells[7].text = "POSITIVO"
                else:
                    row_cells[7].text = "NEGATIVO"
                row_cells[8].text = str(row[7])
                row_cells[9].text = str(row[8])
                
                # Aplica formatação: se o resultado for POSITIVO, bold e vermelho
                for col_index in [7, 8, 9]:
                    cell = row_cells[col_index]
                    if "POSITIVO" in cell.text.upper():
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 0, 0)
            
            for r in table.rows:
                for cell in r.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            tabela_numero += 1

    if count_tables == 0:
        p = doc.add_paragraph("Nenhuma unidade sanitária possui dados para o período selecionado.", style="Heading 2")
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    return doc

def gerar_relatorio(df_atual, df_anterior, periodo_atual_str, periodo_anterior_str, data_emissao, nome_usuario):
    """
    Gera o relatório em Word (landscape) contendo:
      - Cabeçalho institucional com imagem e texto.
      - Período (alinhado à esquerda).
      - Resumo global (alinhado à esquerda), com detalhes para Influenza (quantidade por subtipo).
      - Tabelas específicas por unidade sanitária (na tabela, Influenza mostra apenas POSITIVO/NEGATIVO).
      - Rodapé com o nome do gerador informado pelo usuário.
    """
    doc = Document()

    # Define a orientação para landscape
    from docx.enum.section import WD_ORIENT
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

    # Configura as margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # CABEÇALHO: Insere imagem e texto no cabeçalho real
    header = doc.sections[0].header
    if header.paragraphs:
        header_para = header.paragraphs[0]
    else:
        header_para = header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adiciona a imagem do emblema da República de Moçambique (tamanho reduzido)
    try:
        run_emblem = header_para.add_run()
        run_emblem.add_picture(EMBLEM_PATH, width=Inches(1.0))
    except Exception as e:
        header_para.add_run("\n[Erro ao carregar imagem do Emblema]")
    
    # Adiciona o texto institucional
    header_para.add_run("\nREPÚBLICA DE MOÇAMBIQUE\nMINISTÉRIO DA SAÚDE\nINSTITUTO NACIONAL DE SAÚDE\n")
    
    # CORPO DO DOCUMENTO
    title = doc.add_heading("RELATÓRIO DE RESULTADOS ANALÍTICOS", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("\nPropósito: Vigilância das Infecções Respiratórias Agudas\n").bold = True

    # Período (alinhado à esquerda)
    p_periodo = doc.add_paragraph(f"Período: {periodo_atual_str}", style='Heading 2')
    p_periodo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Resumo global (alinhado à esquerda)
    resumo_texto = gerar_resumo_dinamico(df_atual, df_anterior, periodo_atual_str, periodo_anterior_str)
    p_resumo = doc.add_paragraph(resumo_texto)
    p_resumo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

     # Tabelas por unidade sanitária
    df_atual_sem_data_entrada = df_atual.drop(columns=["Data de entrada"], errors="ignore")
    doc = criar_tabelas_unidades_sanitarias(doc, df_atual_sem_data_entrada)

    # Rodapé
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.add_run(f"\nData de emissão: {data_emissao}")
    footer.add_run(f"\nGerado por: {nome_usuario}")
    footer.add_run(f"\nData e hora do sistema: {CURRENT_DATE}")

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def main():
    st.title("Gerador de Relatórios - Vigilância das Infecções Respiratórias")
    
    st.header("1. Carregar dados")
    uploaded_file = st.file_uploader("Escolha o arquivo Excel ou CSV com os dados", type=['csv', 'xlsx', 'xls'])
    
    if uploaded_file is not None:
        df = carregar_dados(uploaded_file)
        if df is not None:
            st.success("Dados carregados com sucesso!")
            
            # ================= NOVO: Escolha fixa do tipo de data =================
            st.header("2. Filtro de Período (Escolha o tipo de data)")
            coluna_filtro = st.selectbox(
                "Selecione o tipo de data para aplicar o filtro",
                options=[
                    "Data de Testagem SARS",
                    "Data da Testagem FLU",
                    "Data da Testagem RSV",
                    "Data da Colheita"
                ],
                index=3  # por padrão "Data da Colheita"
            )
            # ======================================================================

            col1, col2 = st.columns(2)
            with col1:
                data_inicio = st.date_input("Data inicial", date(2025, 3, 24))
            with col2:
                data_fim = st.date_input("Data final", date(2025, 3, 28))
            
            # Filtra os dados conforme a coluna selecionada
            df_tmp = df.copy()
            df_tmp["_dtFiltro"] = pd.to_datetime(df_tmp.get(coluna_filtro), errors="coerce")
            mask_atual = (df_tmp["_dtFiltro"] >= pd.to_datetime(data_inicio)) & (df_tmp["_dtFiltro"] <= pd.to_datetime(data_fim))
            df_atual = df.loc[mask_atual].copy()

            periodo_atual_str = f"{data_inicio.strftime('%d/%m/%Y')} a {data_fim.strftime('%d/%m/%Y')}"
            st.write(f"Foram encontrados {len(df_atual)} registros no período selecionado (baseado em **{coluna_filtro}**).")
            
            # Semana anterior
            data_inicio_prev = data_inicio - timedelta(days=7)
            data_fim_prev = data_fim - timedelta(days=7)
            mask_anterior = (df_tmp["_dtFiltro"] >= pd.to_datetime(data_inicio_prev)) & (df_tmp["_dtFiltro"] <= pd.to_datetime(data_fim_prev))
            df_anterior = df.loc[mask_anterior].copy()

            periodo_anterior_str = f"{data_inicio_prev.strftime('%d/%m/%Y')} a {data_fim_prev.strftime('%d/%m/%Y')}"
            st.write(f"Foram encontrados {len[df_anterior]} registros no período anterior ({periodo_anterior_str}) (baseado em **{coluna_filtro}**).")
            
            nome_usuario = st.text_input("Nome do Gerador", "Mulungo06")
            data_emissao = st.text_input("Data de Emissão", datetime.now().strftime("%d/%m/%Y"))
            
            if st.button("Gerar Relatório"):
                doc_io = gerar_relatorio(df_atual, df_anterior, periodo_atual_str, periodo_anterior_str, data_emissao, nome_usuario)
                st.download_button(
                    label="📥 Download do Relatório",
                    data=doc_io.getvalue(),
                    file_name=f"relatorio_{data_inicio.strftime('%Y%m%d')}_{data_fim.strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.header("3. Visualização dos Dados (Filtrado)")
            st.dataframe(df_atual)

            # =========================
            # 4. Positividade ao longo do período (por dia)
            # =========================
            st.header("4. Positividade diária por vírus (no período selecionado)")

            # Usar a mesma coluna escolhida no filtro como eixo temporal
            eixo_tempo = coluna_filtro if coluna_filtro in df_atual.columns else "Data da Colheita"
            df_ts = df_atual.copy()
            df_ts["_data"] = pd.to_datetime(df_ts[eixo_tempo], errors="coerce")
            df_ts = df_ts.dropna(subset=["_data"])

            # Marcadores binários de positividade
            df_ts["RSV_Pos"] = df_ts["RSV"].astype(str).str.upper().eq("POSITIVO")
            df_ts["SARS_Pos"] = df_ts["SARS-CoV-2"].astype(str).str.upper().eq("POSITIVO")
            # Influenza (qualquer): já tens "Influenza" com "POSITIVO: ..."; usamos contains
            df_ts["Influenza_Pos"] = df_ts["Influenza"].astype(str).str.upper().str.contains("POSITIVO")

            # Influenza A/B dos flags novos (se existirem); caso não, infere a partir de "Influenza" (fallback)
            if "FluA_Pos" not in df_ts.columns:
                df_ts["FluA_Pos"] = df_ts["Influenza"].astype(str).str.contains(r"A\(", case=False, regex=True)
            if "FluB_Pos" not in df_ts.columns:
                df_ts["FluB_Pos"] = df_ts["Influenza"].astype(str).str.contains(r"\bB\b|\(Victoria\)|\(Yamagata\)", case=False, regex=True)

            # Agregar por dia
            grp = df_ts.groupby(df_ts["_data"].dt.date)
            ts = pd.DataFrame({
                "n_testadas": grp.size()
            })
            for col, name in [
                ("FluA_Pos", "Pos_FluA"),
                ("FluB_Pos", "Pos_FluB"),
                ("RSV_Pos", "Pos_RSV"),
                ("SARS_Pos", "Pos_SARS")
            ]:
                ts[name] = grp[col].sum()

            # Positividade (%)
            for name in ["Pos_FluA", "Pos_FluB", "Pos_RSV", "Pos_SARS"]:
                ts[name.replace("Pos_", "Pct_")] = (ts[name] / ts["n_testadas"] * 100).round(2)

            st.subheader("Positividade (%)")
            st.line_chart(ts[["Pct_FluA", "Pct_FluB", "Pct_RSV", "Pct_SARS"]])

            with st.expander("Ver tabela diária (contagens e percentuais)"):
                st.dataframe(ts.reset_index(names="Data"))

            # =========================
            # 5. Positividade por faixa etária
            # =========================
            st.header("5. Positividade por faixa etária")

            def to_age_num(x):
                v = extrair_valor_idade(x)
                return v if v is not None else float("nan")

            def faixa_etaria(a):
                # Faixas definidas por ti
                if pd.isna(a): return "Idade não informada"
                if a < 2: return "0-2"
                if a < 5: return "2-5"
                if a < 15: return "5-15"
                if a < 50: return "15-50"
                if a < 65: return "50-65"
                return "65+"

            df_age = df_atual.copy()
            df_age["Idade_num"] = df_age["Idade"].apply(to_age_num)
            df_age["Faixa"] = df_age["Idade_num"].apply(faixa_etaria)

            df_age["RSV_Pos"] = df_age["RSV"].astype(str).str.upper().eq("POSITIVO")
            df_age["SARS_Pos"] = df_age["SARS-CoV-2"].astype(str).str.upper().eq("POSITIVO")
            df_age["Influenza_Pos"] = df_age["Influenza"].astype(str).str.upper().str.contains("POSITIVO")

            if "FluA_Pos" not in df_age.columns:
                df_age["FluA_Pos"] = df_age["Influenza"].astype(str).str.contains(r"A\(", case=False, regex=True)
            if "FluB_Pos" not in df_age.columns:
                df_age["FluB_Pos"] = df_age["Influenza"].astype(str).str.contains(r"\bB\b|\(Victoria\)|\(Yamagata\)", case=False, regex=True)

            ag = df_age.groupby("Faixa").agg(
                n=("Faixa", "size"),
                FluA_Pos=("FluA_Pos", "sum"),
                FluB_Pos=("FluB_Pos", "sum"),
                RSV_Pos=("RSV_Pos", "sum"),
                SARS_Pos=("SARS_Pos", "sum")
            ).reset_index()

            for v in ["FluA_Pos", "FluB_Pos", "RSV_Pos", "SARS_Pos"]:
                ag[v.replace("_Pos", "_Pct")] = (ag[v] / ag["n"] * 100).round(2)

            st.subheader("Tabela de positividade por faixa etária (%)")
            st.dataframe(
                ag[["Faixa", "n", "FluA_Pct", "FluB_Pct", "RSV_Pct", "SARS_Pct"]]
                .sort_values("Faixa")
            )

            st.subheader("Gráfico (positividade % por faixa)")
            st.bar_chart(
                ag.set_index("Faixa")[["FluA_Pct", "FluB_Pct", "RSV_Pct", "SARS_Pct"]]
                .sort_index()
            )

            # =========================
            # 6. Mapa por Distrito (opcional)
            # =========================
            st.header("6. Mapa de distribuição por distrito (opcional)")
            st.caption("Carrega um ficheiro GeoJSON de Moçambique com a propriedade 'Distrito' (ex.: limites distritais).")

            geojson_file = st.file_uploader("GeoJSON de Distritos", type=["geojson", "json"], key="geojson_distritos")

            # Contagens por potencial distrito (se existir essa coluna); caso contrário, uso Província/Bairro como fallback
            candidato_distrito = None
            for cand in ["Distrito", "Distrito/Concelho", "Concelho", "Residência/Bairro"]:
                if cand in df_atual.columns:
                    candidato_distrito = cand
                    break

            if geojson_file is not None and candidato_distrito is not None:
                import json, numpy as np, pydeck as pdk

                gj = json.load(geojson_file)
                contagem = (df_atual.groupby(candidato_distrito)
                            .size().reset_index(name="casos"))

                # Projecta contagens para as features do geojson via chave 'Distrito'
                for feat in gj.get("features", []):
                    nome = None
                    # tenta em ordem de propriedades comuns
                    for k in ["Distrito", "district", "name", "NAME_2", "NAME_3", "ADM2_PT", "ADM2_PCODE"]:
                        if "properties" in feat and k in feat["properties"]:
                            nome = str(feat["properties"][k]).strip()
                            break
                    # matching simples (casefold)
                    casos = int(contagem.loc[
                        contagem[candidato_distrito].astype(str).str.casefold() == str(nome).casefold(),
                        "casos"
                    ].sum()) if nome is not None else 0
                    feat.setdefault("properties", {})["casos"] = casos

                st.subheader("Choropleth de casos por distrito")
                layer = pdk.Layer(
                    "GeoJsonLayer",
                    gj,
                    opacity=0.6,
                    stroked=True,
                    get_line_color=[0, 0, 0],
                    get_fill_color="[" 
                                   "min(255, properties.casos * 12), "
                                   "min(255, properties.casos * 6), "
                                   "160]"
                )
                view_state = pdk.ViewState(latitude=-18.9, longitude=35.3, zoom=4.2)  # Moçambique
                st.pydeck_chart(pdk.Deck(layers=[layer], initial_view_state=view_state))
                with st.expander("Ver tabela de contagens por distrito"):
                    st.dataframe(contagem.sort_values("casos", ascending=False))
            else:
                st.info("Sem GeoJSON ou coluna de distrito — a mostrar distribuição tabular/gráfica por Província e por Bairro.")
                if "Província" in df_atual.columns:
                    cnt_prov = df_atual["Província"].value_counts().rename_axis("Província").reset_index(name="Casos")
                    st.subheader("Casos por Província")
                    st.bar_chart(cnt_prov.set_index("Província"))
                    st.dataframe(cnt_prov)
                if "Residência/Bairro" in df_atual.columns:
                    cnt_bairro = (df_atual["Residência/Bairro"]
                                  .value_counts()
                                  .head(25)
                                  .rename_axis("Residência/Bairro")
                                  .reset_index(name="Casos"))
                    st.subheader("Top 25 Bairros por número de casos")
                    st.bar_chart(cnt_bairro.set_index("Residência/Bairro"))
                    st.dataframe(cnt_bairro)

if __name__ == "__main__":
    main()
